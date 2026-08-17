from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Powered Up Peptides Admin Guide.docx"
LOGO = ROOT / "public/assets/brand/logo-wordmark.png"

INK = RGBColor(22, 24, 29)
GOLD = RGBColor(176, 132, 24)
MUTED = RGBColor(95, 99, 108)
PALE_GOLD = "F8F2E2"
LIGHT = "F3F4F6"
WHITE = "FFFFFF"


def font(run, size=11, bold=False, color=INK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for i, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[i]))
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), bold=True)
        font(p.add_run(text[len(bold_lead):]))
    else:
        font(p.add_run(text))
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    font(p.add_run(text))
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.keep_together = True
    font(p.add_run(text))
    return p


def callout(doc, label, text, fill=PALE_GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(label + "  "), size=10.5, bold=True, color=GOLD)
    font(p.add_run(text), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.78)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.38)
section.footer_distance = Inches(0.38)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, GOLD, 18, 10),
    ("Heading 2", 13, GOLD, 14, 7),
    ("Heading 3", 12, INK, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Number"):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.first_line_indent = Inches(-0.25)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

# Running furniture.
hp = section.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(hp.add_run("POWERED UP PEPTIDES  |  ADMIN GUIDE"), size=8.5, bold=True, color=MUTED)
fp = section.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run("Confidential — store administration"), size=8.5, color=MUTED)

# Editorial cover.
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(30)
logo_shape = p.add_run().add_picture(str(LOGO), width=Inches(3.15))
logo_shape._inline.docPr.set("descr", "Powered Up Peptides logo")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
font(p.add_run("STORE ADMINISTRATION GUIDE"), size=10, bold=True, color=GOLD)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
font(p.add_run("Powered Up Peptides"), size=28, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(34)
font(p.add_run("Simple steps for running the store day to day"), size=13, color=MUTED)

callout(doc, "ADMIN LOGIN", "https://powereduppeptides.com/lunar/login")

table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_widths(table, [2700, 6660])
credentials = [
    ("Username / email", "admin@powereduppeptides.com"),
    ("Temporary password", "PoweredUp2026!"),
    ("Prepared", "August 7, 2026"),
]
for idx, (label, value) in enumerate(credentials):
    shade(table.cell(idx, 0), LIGHT)
    shade(table.cell(idx, 1), WHITE)
    p1 = table.cell(idx, 0).paragraphs[0]
    p2 = table.cell(idx, 1).paragraphs[0]
    font(p1.add_run(label), size=10, bold=True, color=MUTED)
    font(p2.add_run(value), size=10.5, bold=idx < 2)

doc.add_paragraph()
callout(doc, "KEEP THIS PRIVATE", "This guide contains the store login. Change the password after the first login and do not share this file publicly.", fill="FDECEC")

doc.add_page_break()

heading(doc, "1. Log in", 1)
numbered(doc, "Go to https://powereduppeptides.com/lunar/login.")
numbered(doc, "Enter the email and password shown on the cover.")
numbered(doc, "To change the password: open Settings → Staff, choose Powered Up Admin, type the new password, then save.")
numbered(doc, "Open the account menu → 2FA Settings and turn on two-factor authentication.")

heading(doc, "2. Your daily order routine", 1)
numbered(doc, "Open Sales → Orders.")
numbered(doc, "Open the newest order and check the customer's name, address, items and total.")
numbered(doc, "Confirm that payment was received outside the website, then change the order to Payment Received.")
numbered(doc, "Pack and ship the order. Only change it to Dispatched after it has actually shipped.")
callout(doc, "IMPORTANT", "The website does not collect money yet. New orders appear as Awaiting Payment. Always confirm payment separately before shipping.", fill="FFF4D6")

heading(doc, "3. Change a price or stock", 1)
bullet(doc, "Open Catalog → Products and choose the product.")
bullet(doc, "Open its variant, then use Pricing to change the price.")
bullet(doc, "Change Stock only when you know the real quantity available.")
bullet(doc, "Save, then open the product on the live website and check your change.")
callout(doc, "Leave these alone", "Do not change the SKU, product type, tax class or technical settings unless the developer asks you to.")

doc.add_page_break()
heading(doc, "4. Change a product image", 1)
bullet(doc, "Open Catalog → Products and choose the product.")
bullet(doc, "Open Media, upload the new image, and make it the main image if needed.")
bullet(doc, "Save, then check the product page on the live website.")

heading(doc, "5. Customers and discount codes", 1)
bullet(doc, "Sales → Customers: look up a customer's details and order history.")
bullet(doc, "Sales → Discounts: create or review a discount code.")
bullet(doc, "Give every discount a clear end date and usage limit.")
bullet(doc, "Test the code in the shopping cart before giving it to customers.")

heading(doc, "6. Lab reports", 1)
body(doc, "Lab reports and certificate PDFs cannot be changed from this admin screen yet. Send the new certificate, batch number, test date, laboratory and purity result to the developer.")
callout(doc, "BAC WATER", "BAC water does not show a peptide purity percentage. It uses sterility and endotoxin testing, so the dash in the Purity column is correct.")

heading(doc, "7. What not to touch", 1)
bullet(doc, "Do not delete products, variants, collections or staff accounts.")
bullet(doc, "Do not change anything under Settings except your password and 2FA.")
bullet(doc, "Do not mark an unpaid order as Payment Received or Dispatched.")
bullet(doc, "If a screen looks technical or unfamiliar, stop and ask before saving.")

doc.add_paragraph()
callout(doc, "WHEN IN DOUBT", "Do not guess. Take a screenshot of the screen and send it to the developer before changing anything.")

doc.core_properties.title = "Powered Up Peptides Admin Guide"
doc.core_properties.subject = "Store administration and operating guide"
doc.core_properties.author = "Powered Up Peptides"
doc.core_properties.keywords = "Lunar, Laravel, ecommerce, admin guide"
doc.save(OUT)
print(OUT)
